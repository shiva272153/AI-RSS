import fitz  # PyMuPDF
from docx import Document
import io


def extract_text_from_pdf(file_bytes):
    """Extract text from PDF file bytes."""
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
        return text.strip()
    except Exception as e:
        print(f"PDF extraction error: {e}")
        return ""


def extract_text_from_docx(file_bytes):
    """Extract text from DOCX file bytes."""
    try:
        doc = Document(io.BytesIO(file_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text.strip()
    except Exception as e:
        print(f"DOCX extraction error: {e}")
        return ""


def extract_text(file_bytes, file_type):
    """Extract text from file based on type."""
    if file_type == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif file_type == 'docx':
        return extract_text_from_docx(file_bytes)
    else:
        return ""
